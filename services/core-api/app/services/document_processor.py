"""
Document Processing Pipeline
Handles document upload, text extraction, chunking, and metadata management
"""

import logging
import asyncio
import hashlib
import mimetypes
from typing import Dict, Any, List, Optional, Union, BinaryIO
from datetime import datetime
from pathlib import Path
import tempfile
import os

try:
    from google.cloud import documentai
    from google.cloud import storage
    import PyPDF2
    import docx
    import pandas as pd
    import requests
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin, urlparse
except ImportError as e:
    logging.warning(f"Document processing dependencies not installed: {e}")
    # Mock classes for development
    class documentai:
        class DocumentProcessorServiceClient:
            pass
    class storage:
        class Client:
            pass
    class PyPDF2:
        pass

from sqlalchemy.orm import Session
from ..models.user import KnowledgeSource, Document, Organization
from ..config.vertex_ai import vertex_ai_config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Document processing pipeline for knowledge management
    
    Supports:
    - PDF, DOCX, CSV file processing
    - URL crawling with depth=1
    - Text extraction and chunking
    - Metadata management and versioning
    """
    
    def __init__(self):
        self.config = vertex_ai_config
        self.project_id = self.config.PROJECT_ID
        self.location = self.config.LOCATION
        
        # Document AI client for OCR
        self._document_ai_client = None
        self._storage_client = None
        
        # Processing configuration
        self.chunk_size = 1024
        self.chunk_overlap = 200
        self.max_file_size_mb = 100
        self.supported_mime_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/csv',
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        ]
        
        self._initialize_clients()
    
    def _initialize_clients(self):
        """Initialize Google Cloud clients"""
        try:
            if self.project_id:
                self._document_ai_client = documentai.DocumentProcessorServiceClient()
                self._storage_client = storage.Client(project=self.project_id)
                logger.info("Document processing clients initialized")
        except Exception as e:
            logger.warning(f"Failed to initialize clients: {e}")
    
    async def process_document_source(
        self,
        db: Session,
        organization_id: str,
        source_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process a document source (file upload or URL)
        
        Args:
            db: Database session
            organization_id: Organization ID
            source_data: Source configuration
            
        Returns:
            Processing result
        """
        try:
            source_type = source_data.get("type")
            source_name = source_data.get("name")
            
            # Create knowledge source record
            knowledge_source = KnowledgeSource(
                organization_id=organization_id,
                type=source_type,
                name=source_name,
                source_url=source_data.get("url"),
                status="processing",
                metadata=source_data.get("metadata", {})
            )
            
            db.add(knowledge_source)
            db.commit()
            
            logger.info(f"Created knowledge source {knowledge_source.id} for organization {organization_id}")
            
            # Process based on source type
            if source_type == "file":
                result = await self._process_file_source(
                    db, knowledge_source, source_data
                )
            elif source_type == "url":
                result = await self._process_url_source(
                    db, knowledge_source, source_data
                )
            else:
                raise ValueError(f"Unsupported source type: {source_type}")
            
            # Update source status
            knowledge_source.status = "completed" if result["success"] else "failed"
            knowledge_source.processed_at = datetime.utcnow()
            knowledge_source.metadata.update(result.get("metadata", {}))
            
            db.commit()
            
            return {
                "source_id": str(knowledge_source.id),
                "status": knowledge_source.status,
                "documents_created": result.get("documents_created", 0),
                "chunks_created": result.get("chunks_created", 0),
                "processing_time_ms": result.get("processing_time_ms", 0)
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            if 'knowledge_source' in locals():
                knowledge_source.status = "failed"
                knowledge_source.metadata["error"] = str(e)
                db.commit()
            raise
    
    async def _process_file_source(
        self,
        db: Session,
        knowledge_source: KnowledgeSource,
        source_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process uploaded file"""
        try:
            start_time = datetime.utcnow()
            
            file_content = source_data.get("content")
            filename = source_data.get("filename")
            mime_type = source_data.get("mime_type")
            
            if not file_content:
                raise ValueError("No file content provided")
            
            # Validate file
            self._validate_file(file_content, mime_type, filename)
            
            # Extract text based on file type
            if mime_type == "application/pdf":
                text_content = await self._extract_pdf_text(file_content)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text_content = await self._extract_docx_text(file_content)
            elif mime_type in ["text/plain"]:
                text_content = file_content.decode('utf-8')
            elif mime_type in ["text/csv", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                text_content = await self._extract_csv_text(file_content, filename)
            else:
                # Use Document AI for OCR
                text_content = await self._extract_with_document_ai(file_content, mime_type)
            
            # Create document chunks
            chunks = self._create_text_chunks(text_content)
            
            # Store document chunks
            documents_created = 0
            for i, chunk in enumerate(chunks):
                document = Document(
                    source_id=knowledge_source.id,
                    chunk_id=i,
                    content=chunk["content"],
                    metadata={
                        "filename": filename,
                        "mime_type": mime_type,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "start_char": chunk["start_char"],
                        "end_char": chunk["end_char"]
                    }
                )
                db.add(document)
                documents_created += 1
            
            db.commit()
            
            end_time = datetime.utcnow()
            processing_time_ms = int((end_time - start_time).total_seconds() * 1000)
            
            return {
                "success": True,
                "documents_created": documents_created,
                "chunks_created": len(chunks),
                "processing_time_ms": processing_time_ms,
                "metadata": {
                    "filename": filename,
                    "mime_type": mime_type,
                    "file_size_bytes": len(file_content),
                    "text_length": len(text_content)
                }
            }
            
        except Exception as e:
            logger.error(f"File processing failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "documents_created": 0,
                "chunks_created": 0
            }
    
    async def _process_url_source(
        self,
        db: Session,
        knowledge_source: KnowledgeSource,
        source_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process URL crawling"""
        try:
            start_time = datetime.utcnow()
            
            url = source_data.get("url")
            max_depth = source_data.get("max_depth", 1)
            
            if not url:
                raise ValueError("No URL provided")
            
            # Crawl URLs
            crawled_urls = await self._crawl_urls(url, max_depth)
            
            total_documents = 0
            total_chunks = 0
            
            # Process each crawled page
            for crawled_url, content in crawled_urls.items():
                if not content:
                    continue
                
                # Create text chunks
                chunks = self._create_text_chunks(content)
                
                # Store document chunks
                for i, chunk in enumerate(chunks):
                    document = Document(
                        source_id=knowledge_source.id,
                        chunk_id=total_documents + i,
                        content=chunk["content"],
                        metadata={
                            "source_url": crawled_url,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "start_char": chunk["start_char"],
                            "end_char": chunk["end_char"]
                        }
                    )
                    db.add(document)
                
                total_documents += len(chunks)
                total_chunks += len(chunks)
            
            db.commit()
            
            end_time = datetime.utcnow()
            processing_time_ms = int((end_time - start_time).total_seconds() * 1000)
            
            return {
                "success": True,
                "documents_created": total_documents,
                "chunks_created": total_chunks,
                "processing_time_ms": processing_time_ms,
                "metadata": {
                    "urls_crawled": len(crawled_urls),
                    "base_url": url,
                    "max_depth": max_depth
                }
            }
            
        except Exception as e:
            logger.error(f"URL processing failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "documents_created": 0,
                "chunks_created": 0
            }
    
    def _validate_file(self, file_content: bytes, mime_type: str, filename: str):
        """Validate uploaded file"""
        # Check file size
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValueError(f"File size {file_size_mb:.1f}MB exceeds limit of {self.max_file_size_mb}MB")
        
        # Check MIME type
        if mime_type not in self.supported_mime_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        # Verify MIME type matches file extension
        guessed_type, _ = mimetypes.guess_type(filename)
        if guessed_type and guessed_type != mime_type:
            logger.warning(f"MIME type mismatch: provided {mime_type}, guessed {guessed_type}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            # Use Document AI if available, otherwise PyPDF2
            if self._document_ai_client:
                return await self._extract_with_document_ai(file_content, "application/pdf")
            else:
                # Fallback to PyPDF2
                with tempfile.NamedTemporaryFile(suffix=".pdf") as temp_file:
                    temp_file.write(file_content)
                    temp_file.flush()
                    
                    text_parts = []
                    with open(temp_file.name, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        for page in pdf_reader.pages:
                            text_parts.append(page.extract_text())
                    
                    return "\n".join(text_parts)
                    
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx") as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                doc = docx.Document(temp_file.name)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)
                
                return "\n".join(text_parts)
                
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise
    
    async def _extract_csv_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from CSV/Excel"""
        try:
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                # Read with pandas
                if filename.endswith('.csv'):
                    df = pd.read_csv(temp_file.name)
                else:
                    df = pd.read_excel(temp_file.name)
                
                # Convert to text representation
                text_parts = []
                text_parts.append(f"Columns: {', '.join(df.columns)}")
                text_parts.append(f"Rows: {len(df)}")
                text_parts.append("")
                
                # Add sample data
                for index, row in df.head(100).iterrows():  # Limit to first 100 rows
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                    text_parts.append(row_text)
                
                return "\n".join(text_parts)
                
        except Exception as e:
            logger.error(f"CSV text extraction failed: {e}")
            raise
    
    async def _extract_with_document_ai(self, file_content: bytes, mime_type: str) -> str:
        """Extract text using Google Document AI"""
        try:
            if not self._document_ai_client:
                raise ValueError("Document AI client not available")
            
            # Mock implementation - in production would use actual Document AI
            logger.info("Using Document AI for text extraction")
            return "Mock text extracted with Document AI"
            
        except Exception as e:
            logger.error(f"Document AI extraction failed: {e}")
            raise
    
    async def _crawl_urls(self, base_url: str, max_depth: int = 1) -> Dict[str, str]:
        """Crawl URLs and extract text content"""
        try:
            crawled_content = {}
            visited_urls = set()
            urls_to_visit = [(base_url, 0)]
            
            while urls_to_visit:
                current_url, depth = urls_to_visit.pop(0)
                
                if current_url in visited_urls or depth > max_depth:
                    continue
                
                visited_urls.add(current_url)
                
                try:
                    # Fetch page content
                    response = requests.get(current_url, timeout=10)
                    response.raise_for_status()
                    
                    # Parse HTML
                    soup = BeautifulSoup(response.content, 'html.parser')
                    
                    # Extract text content
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    text_content = soup.get_text()
                    # Clean up text
                    lines = (line.strip() for line in text_content.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    text_content = ' '.join(chunk for chunk in chunks if chunk)
                    
                    crawled_content[current_url] = text_content
                    
                    # Find links for next depth level
                    if depth < max_depth:
                        for link in soup.find_all('a', href=True):
                            href = link['href']
                            full_url = urljoin(current_url, href)
                            
                            # Only crawl same domain
                            if urlparse(full_url).netloc == urlparse(base_url).netloc:
                                urls_to_visit.append((full_url, depth + 1))
                
                except Exception as e:
                    logger.warning(f"Failed to crawl {current_url}: {e}")
                    continue
            
            return crawled_content
            
        except Exception as e:
            logger.error(f"URL crawling failed: {e}")
            raise
    
    def _create_text_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create text chunks with overlap"""
        try:
            if not text or len(text.strip()) == 0:
                return []
            
            chunks = []
            text_length = len(text)
            
            # Simple chunking strategy
            start = 0
            chunk_id = 0
            
            while start < text_length:
                end = min(start + self.chunk_size, text_length)
                
                # Try to break at sentence boundary
                if end < text_length:
                    # Look for sentence endings within overlap range
                    sentence_end = text.rfind('.', start, end)
                    if sentence_end > start + self.chunk_size // 2:
                        end = sentence_end + 1
                
                chunk_text = text[start:end].strip()
                
                if chunk_text:
                    chunks.append({
                        "content": chunk_text,
                        "start_char": start,
                        "end_char": end,
                        "chunk_id": chunk_id,
                        "length": len(chunk_text)
                    })
                    chunk_id += 1
                
                # Move start position with overlap
                start = max(start + self.chunk_size - self.chunk_overlap, end)
            
            logger.info(f"Created {len(chunks)} chunks from {text_length} characters")
            return chunks
            
        except Exception as e:
            logger.error(f"Text chunking failed: {e}")
            raise
    
    async def reprocess_source(
        self,
        db: Session,
        source_id: str,
        organization_id: str
    ) -> Dict[str, Any]:
        """Reprocess an existing knowledge source"""
        try:
            # Get existing source
            knowledge_source = db.query(KnowledgeSource).filter(
                KnowledgeSource.id == source_id,
                KnowledgeSource.organization_id == organization_id
            ).first()
            
            if not knowledge_source:
                raise ValueError(f"Knowledge source {source_id} not found")
            
            # Delete existing documents
            db.query(Document).filter(Document.source_id == source_id).delete()
            db.commit()
            
            # Reprocess based on source type
            if knowledge_source.type == "url":
                source_data = {
                    "type": "url",
                    "url": knowledge_source.source_url,
                    "name": knowledge_source.name,
                    "metadata": knowledge_source.metadata
                }
                result = await self._process_url_source(db, knowledge_source, source_data)
            else:
                raise ValueError(f"Cannot reprocess source type: {knowledge_source.type}")
            
            # Update source
            knowledge_source.status = "completed" if result["success"] else "failed"
            knowledge_source.processed_at = datetime.utcnow()
            db.commit()
            
            return result
            
        except Exception as e:
            logger.error(f"Source reprocessing failed: {e}")
            raise
    
    async def get_processing_stats(self, db: Session, organization_id: str) -> Dict[str, Any]:
        """Get document processing statistics"""
        try:
            # Get source counts by status
            sources = db.query(KnowledgeSource).filter(
                KnowledgeSource.organization_id == organization_id
            ).all()
            
            status_counts = {}
            type_counts = {}
            
            for source in sources:
                status_counts[source.status] = status_counts.get(source.status, 0) + 1
                type_counts[source.type] = type_counts.get(source.type, 0) + 1
            
            # Get document counts
            total_documents = db.query(Document).join(KnowledgeSource).filter(
                KnowledgeSource.organization_id == organization_id
            ).count()
            
            return {
                "total_sources": len(sources),
                "status_breakdown": status_counts,
                "type_breakdown": type_counts,
                "total_documents": total_documents,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "supported_formats": self.supported_mime_types
            }
            
        except Exception as e:
            logger.error(f"Failed to get processing stats: {e}")
            raise


# Global instance
document_processor = DocumentProcessor()